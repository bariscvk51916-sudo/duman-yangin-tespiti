from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

GITHUB_LINK = "https://github.com/bariscvk51916-sudo/duman-yangin-tespiti"
OUTPUT = r"c:\Users\Baris\Desktop\DİJİTAL GÖRÜNTÜ CÇÖZÜMLEME FİNAL\Baris_Cevik_Donem_Projesi_Raporu.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# Kapak
kapak = doc.add_paragraph()
kapak.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = kapak.add_run("MERSİN ÜNİVERSİTESİ\n\n")
r.bold = True
r.font.size = Pt(14)

r2 = kapak.add_run("Dijital Görüntü Çözümleme\nDönem Projesi Raporu\n\n")
r2.bold = True
r2.font.size = Pt(14)

r3 = kapak.add_run("Duman ve Yangın Tespit Sistemi")
r3.bold = True
r3.font.size = Pt(13)

doc.add_paragraph()

bilgi = doc.add_paragraph()
bilgi.alignment = WD_ALIGN_PARAGRAPH.CENTER
bilgi.add_run("Barış Çevik\n")
bilgi.add_run("Öğrenci No: 22430070008")

doc.add_page_break()

def baslik(metin):
    p = doc.add_paragraph()
    run = p.add_run(metin)
    run.bold = True

def yazi(metin):
    doc.add_paragraph(metin)

baslik("Giriş")

yazi(
    "Bu rapor dijital görüntü çözümleme dersi için hazırladığım dönem projesini anlatıyor. "
    "Projede temel amaç bir fotoğraf verildiğinde görüntüde duman veya yangın olup olmadığını "
    "bulmak. Bunu yaparken derste gördüğümüz nesne tespiti mantığını kullandım."
)

yazi(
    "Hoca derste plaka tanıma örneği vermişti, benzer bir sistem kurmak istedim ama aynı konuyu "
    "seçmek istemedim. Yangın ve duman konusu hem projeye uyuyordu hem de gerçek hayatta işe "
    "yarayabilecek bir şey gibi geldi."
)

baslik("Projenin Amacı")

yazi(
    "Kısaca anlatmak gerekirse kullanıcı bilgisayardan bir görsel yüklüyor, program o görseli "
    "inceliyor ve yangın ya da duman varsa ekranda gösteriyor. Sadece kod çalışsın diye değil, "
    "basit bir arayüz de olsun istedim çünkü hocanın istediği şey buydu."
)

baslik("Kullandığım Programlar")

yazi(
    "YOLO modelini ultralytics kütüphanesi ile kullandım. Görüntü okuma tarafında opencv var. "
    "Web arayüzü için streamlit tercih ettim, flask'a göre daha hızlı kuruluyordu. Veri setini "
    "roboflow sitesinden indirdim. Kendi bilgisayarımda güçlü bir ekran kartı olmadığı için "
    "model eğitimini google colab üzerinden yaptım."
)

baslik("Veri Seti Hakkında")

yazi(
    "Eğitim için duman ve yangın içeren hazır etiketli görseller kullandım. Veriyi yolov8 "
    "formatında indirdim. Train ve valid diye ayrılmış halde geldi zaten. Hoca en az 200 görüntü "
    "dediği için buna dikkat ettim."
)

yazi(
    "Görüntü dosyaları çok büyük olduğu için hepsini githuba yüklemedim. Repoda kodlar ve "
    "açıklamalar var, veri kısmı ayrı duruyor."
)

baslik("Model Eğitimi")

yazi(
    "Eğitimi colabda yaptım. Önce roboflowdan indirdiğim zip dosyasını yükledim, sonra notebook "
    "üzerinden yolo modelini çalıştırdım. Eğitim biraz zaman aldı ama gpu olduğu için bilgisayarıma "
    "göre daha rahat bitti."
)

yazi(
    "Eğitim bitince best.pt dosyasını indirip projenin models klasörüne koydum. Bu dosya olmadan "
    "arayüz çalışıyor ama tahmin yapamıyor, o yüzden önemli."
)

baslik("Metrikler")

yazi(
    "Modeli değerlendirirken mAP, precision ve recall değerlerine baktım. Ayrıca eğitim sırasında "
    "oluşan grafikleri de rapora ekledim. results grafiği ve confusion matrix bunlardan. Bu sayılar "
    "modelin ne kadar doğru tahmin yaptığını gösteriyor."
)

baslik("Arayüz")

yazi(
    "Streamlit ile basit bir sayfa yaptım. Kullanıcı görsel seçiyor, ben de arka planda modeli "
    "çalıştırıyorum. Sonuç varsa kutucuklar çiziliyor ve hangi sınıf olduğu yazıyor."
)

yazi(
    "Çalıştırmak için terminalde proje klasörüne girip streamlit run app/streamlit_app.py "
    "yazmak yeterli. Tarayıcı kendisi açılıyor."
)

baslik("Dosya Yapısı")

yazi(
    "app klasöründe arayüz, src klasöründe tespit kodu var. models klasörüne eğitilmiş dosya "
    "geliyor. notebooks klasöründe colab dosyası duruyor. data klasörüne de eğitim görselleri "
    "konuluyor."
)

baslik("Sonuç")

yazi(
    "Proje boyunca veri bulma, model eğitme ve arayüz yapma adımlarını tamamladım. Ortaya "
    "fotoğraf yükleyip yangın/duman kontrolü yapabilen bir sistem çıktı. Derste anlatılan konuları "
    "pratikte uygulamak için iyi bir fırsat oldu."
)

yazi(
    "İleride bu sisteme canlı kamera bağlanabilir veya alarm eklenebilir diye düşünüyorum ama "
    "şimdilik ders kapsamında bu hali yeterli."
)

doc.add_paragraph()
doc.add_paragraph()

# GitHub linki en altta
alt = doc.add_paragraph()
alt.alignment = WD_ALIGN_PARAGRAPH.CENTER
alt.add_run("Proje GitHub adresi:\n")
alt.add_run(GITHUB_LINK)

doc.save(OUTPUT)
print("Rapor olusturuldu:", OUTPUT)
